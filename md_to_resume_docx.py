# -*- coding: utf-8 -*-
"""Convert resume_optimized_zh.md to a styled .docx with minimal content changes."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Pt
from docx.oxml.ns import qn


def set_run_east_asia(run, font_name: str = "微软雅黑") -> None:
    run.font.name = font_name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), font_name)


def add_body_paragraph(
    doc,
    text: str,
    bold: bool = False,
    size: float = 10.5,
    align=None,
) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.15
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    set_run_east_asia(run)


def add_heading_custom(
    doc, text: str, level: int, space_before: float = 6, space_after: float = 3
) -> None:
    sizes = {1: 22, 2: 14, 3: 12, 4: 11}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = level >= 2
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(sizes.get(level, 11))
    set_run_east_asia(run)


def add_bullet(doc, text: str, size: float = 10.5) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(text)
    run.font.size = Pt(size)
    set_run_east_asia(run)


def main() -> None:
    base = Path(__file__).resolve().parent
    md_path = base / "resume_optimized_zh.md"
    out_dir = base / "output" / "doc"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "resume_optimized_zh.docx"

    raw = md_path.read_text(encoding="utf-8")
    lines = raw.splitlines()

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)

    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(10.5)
    if style.element.rPr is not None and style.element.rPr.rFonts is not None:
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    i = 0
    after_title = False  # 姓名下方、首个二级标题前的联系信息行，居中展示
    while i < len(lines):
        line = lines[i].rstrip("\r\n")  # 仅去掉换行，保留原文行尾空格

        if line.strip() == "":
            i += 1
            continue

        if line.startswith("# "):
            add_heading_custom(doc, line[2:].strip(), 1, space_before=0, space_after=8)
            after_title = True
            i += 1
            continue
        if line.startswith("## "):
            after_title = False
            add_heading_custom(doc, line[3:].strip(), 2, space_before=10, space_after=4)
            i += 1
            continue
        if line.startswith("### "):
            add_heading_custom(doc, line[4:].strip(), 3, space_before=8, space_after=3)
            i += 1
            continue
        if line.startswith("#### "):
            add_heading_custom(doc, line[5:].strip(), 4, space_before=4, space_after=2)
            i += 1
            continue

        if line.startswith("- "):
            add_bullet(doc, line[2:].strip())
            i += 1
            continue

        if after_title:
            add_body_paragraph(doc, line, align=WD_ALIGN_PARAGRAPH.CENTER)
        else:
            add_body_paragraph(doc, line)
        i += 1

    doc.save(out_path)
    print(f"Wrote: {out_path}")


if __name__ == "__main__":
    main()
